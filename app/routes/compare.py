"""Document comparison routes — fast single-LLM-call approach + DOCX export."""
from __future__ import annotations

import io
import os
import re
from datetime import datetime

from fastapi import APIRouter
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

import logging
import requests

from app.config import TABLE_DOCUMENTS, TABLE_DOCUMENT_SECTIONS
from app.services.llm_service import call_llm_chat, _get_token
from app.services.uc_repository import execute_sql

logger = logging.getLogger(__name__)

DIFF_MODEL = "databricks-claude-opus-4-8"

_page_text_cache: dict[str, dict[int, str]] = {}


def _get_page_text(document_id: str, page_number: int) -> str:
    """Get the full text of a specific page directly from the PDF."""
    from app.services.uc_repository import download_file_from_volume

    if document_id in _page_text_cache and page_number in _page_text_cache[document_id]:
        return _page_text_cache[document_id][page_number]

    doc_rows = execute_sql(
        f"SELECT uc_volume_path FROM {TABLE_DOCUMENTS} WHERE document_id = '{document_id}'"
    )
    if not doc_rows:
        return ""

    uc_path = doc_rows[0].get("uc_volume_path", "")
    file_bytes = download_file_from_volume(uc_path)
    if not file_bytes:
        return ""

    try:
        import fitz
        pdf = fitz.open(stream=file_bytes, filetype="pdf")
        if document_id not in _page_text_cache:
            _page_text_cache[document_id] = {}
        for i in range(len(pdf)):
            text = pdf[i].get_text("text").strip()
            _page_text_cache[document_id][i + 1] = text
        pdf.close()
        if len(_page_text_cache) > 10:
            oldest = next(iter(_page_text_cache))
            del _page_text_cache[oldest]
        return _page_text_cache[document_id].get(page_number, "")
    except Exception as e:
        logger.error("Failed to extract page text: %s", e)
        return ""


def _call_diff_llm(messages: list[dict]) -> str:
    """Call the expensive high-accuracy model for page diffs."""
    from app.config import DATABRICKS_HOST
    host = DATABRICKS_HOST.rstrip("/") if DATABRICKS_HOST else ""
    if not host:
        try:
            from databricks.sdk import WorkspaceClient
            host = WorkspaceClient().config.host.rstrip("/")
        except Exception:
            return ""

    url = f"{host}/serving-endpoints/{DIFF_MODEL}/invocations"
    token = _get_token()
    try:
        resp = requests.post(
            url,
            headers={"Authorization": f"Bearer {token}", "Content-Type": "application/json"},
            json={"messages": messages, "temperature": 0.0, "max_tokens": 1024},
            timeout=60,
        )
        resp.raise_for_status()
        return resp.json()["choices"][0]["message"]["content"]
    except Exception as e:
        logger.warning("Diff LLM (%s) failed: %s — falling back to default", DIFF_MODEL, e)
        return call_llm_chat(messages, temperature=0.0)

router = APIRouter(prefix="/api/compare", tags=["compare"])

_APP_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
_LOGO_PATH = os.path.join(_APP_DIR, "frontend", "dist", "images", "element-logo.png")


class CompareRequest(BaseModel):
    document_ids: list[str]


class DocxRequest(BaseModel):
    document_ids: list[str]
    conclusion: str
    include_page_diffs: bool = True


@router.post("")
async def run_comparison(req: CompareRequest):
    """Compare 2+ documents — single SQL query + single LLM call for speed."""
    if len(req.document_ids) < 2:
        return {"error": "Need at least 2 documents to compare"}

    id_list = ", ".join(f"'{d}'" for d in req.document_ids)

    docs = execute_sql(f"""
        SELECT document_id, original_file_name, spec_number, issue_year, status
        FROM {TABLE_DOCUMENTS}
        WHERE document_id IN ({id_list})
        ORDER BY issue_year
    """)

    if len(docs) < 2:
        return {"error": "Documents not found"}

    sections = execute_sql(f"""
        SELECT document_id, section_title, SUBSTRING(section_text, 1, 2500) as section_text,
               start_page, section_number
        FROM {TABLE_DOCUMENT_SECTIONS}
        WHERE document_id IN ({id_list})
        ORDER BY document_id, section_number
        LIMIT 40
    """)

    if not sections:
        return {"error": "No sections found — documents may still be processing"}

    doc_map = {d["document_id"]: d for d in docs}

    context_parts = []
    for s in sections:
        doc = doc_map.get(s["document_id"], {})
        context_parts.append(
            f"[{doc.get('original_file_name','?')} | Year: {doc.get('issue_year','?')} | Page {s.get('start_page','?')} | Section {s.get('section_number','?')}: {s.get('section_title','?')}]\n"
            f"{s.get('section_text','')}\n"
        )
    context = "\n---\n".join(context_parts)

    doc_list = ", ".join(f"{d.get('original_file_name')} ({d.get('issue_year')})" for d in docs)

    messages = [
        {"role": "system", "content": """You are an expert engineering specification analyst. Compare document versions and output a structured analysis.

RESPOND IN THIS EXACT FORMAT:

## Overall Conclusion
[2-3 sentences summarizing the key evolution between versions]

## Key Differences

### 1. [Topic/Section Name]
- **Change type**: added/removed/modified
- **Risk level**: low/medium/high
- **Summary**: [What changed and why it matters]

### 2. [Next Topic]
...

List the most significant 5-8 differences. Be specific — quote exact requirement changes where possible. Note significance for engineering/quality teams."""},
        {"role": "user", "content": f"DOCUMENTS: {doc_list}\n\nCONTEXT:\n{context}\n\nProvide a structured comparison of these specification versions."},
    ]

    answer = call_llm_chat(messages)

    if not answer:
        return {"error": "LLM comparison failed — please try again"}

    return {
        "conclusion": answer,
        "pairwise_comparisons": [],
        "documents": docs,
    }


class PageDiffRequest(BaseModel):
    doc_id_old: str
    doc_id_new: str
    page_number: int


@router.post("/page-diff")
async def get_page_diff(req: PageDiffRequest):
    """Get a precise LLM summary of actual differences on this page between two docs.
    
    Uses the raw PDF page text (not section titles) to avoid text-block extraction artifacts.
    """
    from app.services.uc_repository import download_file_from_volume

    old_page_text = _get_page_text(req.doc_id_old, req.page_number)
    new_page_text = _get_page_text(req.doc_id_new, req.page_number)

    if not old_page_text and not new_page_text:
        return {"page": req.page_number, "summary": "No text found on this page.", "changes": []}

    old_text = old_page_text or "(page not found)"
    new_text = new_page_text or "(page not found)"

    messages = [
        {"role": "system", "content": """You are a precision document comparison tool. Your job is to find ONLY REAL differences between two versions of a specification page.

ABSOLUTE RULES — VIOLATION MEANS FAILURE:
1. NEVER report something as "changed" if the same text appears in both OLD and NEW versions.
2. Formatting differences (line breaks, spacing, capitalization style) are NOT content changes — ignore them.
3. If a title reads "METALLIC MATERIALS" in OLD and "METALLIC MATERIALS" in NEW — that is NOT a change. Do not report it.
4. Only report a difference if the actual WORDS or NUMBERS are different between OLD and NEW.
5. Dates, version numbers, added/removed bullet points, new paragraphs — these are real changes.
6. If the two texts are essentially the same content, say so.

OUTPUT FORMAT (follow exactly):
SUMMARY: [One factual sentence about what genuinely differs, or "No meaningful changes on this page."]
CHANGES:
- [what specifically changed — quote the OLD vs NEW text] | [modified/added/removed]

If no real differences exist:
SUMMARY: No meaningful changes on this page.
CHANGES:
(none)"""},
        {"role": "user", "content": f"Compare these two versions of page {req.page_number}. Find ONLY genuine content differences (not formatting).\n\n=== OLD VERSION ===\n{old_text}\n\n=== NEW VERSION ===\n{new_text}"},
    ]

    answer = _call_diff_llm(messages)
    if not answer:
        return {"page": req.page_number, "summary": "Could not analyze this page.", "changes": []}

    lines = answer.strip().split("\n")
    summary = ""
    changes = []
    for line in lines:
        if line.startswith("SUMMARY:"):
            summary = line[8:].strip()
        elif line.startswith("- "):
            change_text = line[2:].strip()
            if change_text and change_text != "(none)":
                changes.append(change_text)

    if not summary:
        summary = answer[:150]

    return {"page": req.page_number, "summary": summary, "changes": changes}


@router.post("/download-docx")
async def download_comparison_docx(req: DocxRequest):
    """Generate a branded DOCX report from comparison results including per-page diffs."""
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    id_list = ", ".join(f"'{d}'" for d in req.document_ids)
    docs = execute_sql(f"""
        SELECT document_id, original_file_name, spec_number, issue_year, status, page_count
        FROM {TABLE_DOCUMENTS}
        WHERE document_id IN ({id_list})
        ORDER BY issue_year
    """)

    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if os.path.exists(_LOGO_PATH):
        run = header_para.add_run()
        run.add_picture(_LOGO_PATH, width=Inches(1.8))

    header_right = header_para.add_run("\tConfidential")
    header_right.font.size = Pt(8)
    header_right.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(f"Generated by Element Spec Intelligence Platform • {datetime.now().strftime('%d %B %Y')}")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    title = doc.add_heading("Specification Comparison Report", level=0)
    title.runs[0].font.color.rgb = RGBColor(0x00, 0x2B, 0x5C)

    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Document"
    hdr_cells[1].text = "Spec Number"
    hdr_cells[2].text = "Year"
    hdr_cells[3].text = "Status"
    for d in docs:
        row = table.add_row().cells
        row[0].text = d.get("original_file_name", "")
        row[1].text = d.get("spec_number", "")
        row[2].text = str(d.get("issue_year", ""))
        row[3].text = d.get("status", "")

    doc.add_paragraph()

    # --- AI Analysis / Overall Conclusion ---
    _render_markdown_to_docx(doc, req.conclusion)

    # --- Per-Page Section Differences ---
    if req.include_page_diffs and len(req.document_ids) >= 2:
        doc.add_page_break()
        h = doc.add_heading("Page-by-Page Detailed Differences", level=1)
        h.runs[0].font.color.rgb = RGBColor(0x00, 0x2B, 0x5C)

        intro = doc.add_paragraph()
        intro_run = intro.add_run(
            "The following section provides a page-by-page comparison of every change between the OLD and NEW documents. "
            "Each entry identifies exactly what was modified, added, or removed on that page."
        )
        intro_run.font.size = Pt(10)
        intro_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        doc.add_paragraph()

        old_doc_id = req.document_ids[0]
        new_doc_id = req.document_ids[1]

        old_doc_info = next((d for d in docs if d["document_id"] == old_doc_id), None)
        new_doc_info = next((d for d in docs if d["document_id"] == new_doc_id), None)

        max_pages = max(
            int(old_doc_info.get("page_count", 0)) if old_doc_info else 0,
            int(new_doc_info.get("page_count", 0)) if new_doc_info else 0,
        )
        max_pages = min(max_pages, 80)  # cap to avoid timeout

        pages_with_changes = []
        pages_no_changes = []

        for page_num in range(1, max_pages + 1):
            old_text = _get_page_text(old_doc_id, page_num)
            new_text = _get_page_text(new_doc_id, page_num)

            if not old_text and not new_text:
                continue

            if not old_text:
                pages_with_changes.append({
                    "page": page_num,
                    "summary": f"Page {page_num} is newly added in the NEW version.",
                    "changes": [f"Entire page added (not present in OLD document)"]
                })
                continue
            if not new_text:
                pages_with_changes.append({
                    "page": page_num,
                    "summary": f"Page {page_num} was removed in the NEW version.",
                    "changes": [f"Entire page removed (not present in NEW document)"]
                })
                continue

            if old_text.strip() == new_text.strip():
                pages_no_changes.append(page_num)
                continue

            messages = [
                {"role": "system", "content": """You are a precision document comparison tool. Find ONLY REAL content differences between two page versions.

RULES:
1. NEVER report formatting-only changes (spacing, line breaks, capitalization style).
2. Only report genuine word/number/sentence changes.
3. Quote the exact OLD text vs NEW text for each difference.
4. Be comprehensive — list EVERY real difference on this page.

OUTPUT FORMAT:
SUMMARY: [One sentence describing what changed on this page]
CHANGES:
- "OLD text" vs "NEW text" | [modified/added/removed]"""},
                {"role": "user", "content": f"Compare page {page_num}.\n\n=== OLD ===\n{old_text[:4000]}\n\n=== NEW ===\n{new_text[:4000]}"},
            ]

            answer = _call_diff_llm(messages)
            if not answer:
                continue

            lines = answer.strip().split("\n")
            summary = ""
            changes = []
            for line in lines:
                if line.startswith("SUMMARY:"):
                    summary = line[8:].strip()
                elif line.startswith("- "):
                    change_text = line[2:].strip()
                    if change_text and change_text != "(none)":
                        changes.append(change_text)

            if summary and "no meaningful" not in summary.lower() and "no change" not in summary.lower():
                pages_with_changes.append({"page": page_num, "summary": summary, "changes": changes})
            else:
                pages_no_changes.append(page_num)

        # Write pages with changes
        for entry in pages_with_changes:
            page_heading = doc.add_heading(f"Page {entry['page']}", level=2)
            page_heading.runs[0].font.color.rgb = RGBColor(0x00, 0x4D, 0x8C)

            summary_para = doc.add_paragraph()
            summary_run = summary_para.add_run(entry["summary"])
            summary_run.font.size = Pt(10)
            summary_run.bold = True

            if entry.get("changes"):
                for change in entry["changes"]:
                    para = doc.add_paragraph(style="List Bullet")
                    _add_bold_text(para, change)

            doc.add_paragraph()

        # Summary of unchanged pages
        if pages_no_changes:
            doc.add_paragraph()
            unchanged_heading = doc.add_heading("Unchanged Pages", level=2)
            unchanged_heading.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            unchanged_para = doc.add_paragraph()
            page_ranges = _format_page_ranges(pages_no_changes)
            unch_run = unchanged_para.add_run(f"The following pages have no meaningful content changes: {page_ranges}")
            unch_run.font.size = Pt(10)
            unch_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # --- Disclaimer ---
    doc.add_paragraph()
    disclaimer = doc.add_paragraph()
    disclaimer_run = disclaimer.add_run(
        "Disclaimer: This comparison was generated by AI and should be reviewed by qualified engineers. "
        "Always refer to the original specification documents for authoritative information."
    )
    disclaimer_run.font.size = Pt(8)
    disclaimer_run.font.italic = True
    disclaimer_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    spec_num = docs[0].get("spec_number", "spec") if docs else "spec"
    filename = f"Comparison_Report_{spec_num}_{datetime.now().strftime('%Y%m%d')}.docx"

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _format_page_ranges(pages: list[int]) -> str:
    """Format a list of page numbers into compact ranges like '1-5, 8, 10-12'."""
    if not pages:
        return ""
    pages = sorted(pages)
    ranges = []
    start = pages[0]
    end = pages[0]
    for p in pages[1:]:
        if p == end + 1:
            end = p
        else:
            ranges.append(f"{start}" if start == end else f"{start}-{end}")
            start = p
            end = p
    ranges.append(f"{start}" if start == end else f"{start}-{end}")
    return ", ".join(ranges)


def _render_markdown_to_docx(doc, text: str):
    """Convert markdown-formatted text to DOCX paragraphs with formatting."""
    from docx.shared import Pt, RGBColor

    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        if line.startswith("## "):
            heading = doc.add_heading(line[3:].strip(), level=1)
            heading.runs[0].font.color.rgb = RGBColor(0x00, 0x2B, 0x5C)
        elif line.startswith("### "):
            heading = doc.add_heading(line[4:].strip(), level=2)
            heading.runs[0].font.color.rgb = RGBColor(0x00, 0x4D, 0x8C)
        elif line.startswith("- **") or line.startswith("* **"):
            para = doc.add_paragraph(style="List Bullet")
            _add_bold_text(para, line[2:].strip())
        elif line.startswith("- ") or line.startswith("* "):
            para = doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line.strip() == "":
            pass
        else:
            para = doc.add_paragraph()
            _add_bold_text(para, line.strip())

        i += 1


def _add_bold_text(para, text: str):
    """Parse **bold** markers in text and add formatted runs."""
    from docx.shared import Pt

    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(10)
        else:
            run = para.add_run(part)
            run.font.size = Pt(10)
